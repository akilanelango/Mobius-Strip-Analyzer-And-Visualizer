import os
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
from mobius_strip import MobiusStrip
import gc

# File paths
testcase_path = "./mobius_testcases.txt"
output_doc_path = "./Mobius_Strip_Results.docx"
plot_folder = "./mobius_plots"

# Create folder to store plot images
os.makedirs(plot_folder, exist_ok=True)

# Load test cases, ignoring comment lines
with open(testcase_path, "r") as file:
    lines = file.readlines()

# Initialize Word document
doc = Document()
doc.add_heading("Möbius Strip Test Results", 0)

plot_index = 1

# Process each test case
for line in lines:
    line = line.strip()
    if not line or line.startswith("#"):
        continue  # Skip empty lines and comments

    try:
        R, w, n = map(float, line.split())
        n = int(n)

        # Generate and analyze Möbius strip
        mobius = MobiusStrip(R=R, w=w, n=n)
        (area, edge_length), fig = mobius.analyze()

        # Save plot
        plot_filename = os.path.join(plot_folder, f"plot_{plot_index}.jpg")
        fig.savefig(plot_filename, format="jpg", dpi=100)
        plt.close(fig)
        gc.collect()

        # Add results to document
        doc.add_heading(f"Test Case {plot_index}", level=1)
        doc.add_paragraph(f"Parameters: R = {R}, w = {w}, n = {n}")
        doc.add_paragraph(f"Surface Area ≈ {area:.4f}")
        doc.add_paragraph(f"Edge Length ≈ {edge_length:.4f}")
        doc.add_picture(plot_filename, width=Inches(5))

        plot_index += 1

    except Exception as e:
        doc.add_heading(f"Test Case {plot_index} - Error", level=1)
        doc.add_paragraph(f"Error processing line: {line}")
        doc.add_paragraph(str(e))
        plot_index += 1

# Save final Word document
doc.save(output_doc_path)

print(f"Test results saved to: {output_doc_path}")

